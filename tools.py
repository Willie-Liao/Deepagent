"""
Custom tools for IB MYP PHE Criterion C Prompt Writer.
"""

import os
import requests
from typing import Literal, Optional
from pathlib import Path


# Workspace root for file operations (matches agent.py WORKSPACE_ROOT)
WORKSPACE_ROOT = Path(__file__).parent / "workspace"


def get_criterion_reference(
    year: Literal["Year 1", "Year 3", "Year 5", "year 1", "year 3", "year 5", "1", "3", "5"],
    criterion: Literal["A", "B", "C", "D"] = "C"
) -> str:
    """
    Load achievement level descriptors from the MYP assessment criteria reference.
    
    Args:
        year: MYP year level (Year 1, Year 3, or Year 5)
        criterion: Assessment criterion (A, B, C, or D). Default is C.
    
    Returns:
        The achievement level descriptors for the specified year and criterion.
    """
    # Normalize year input
    year_map = {
        "1": "Year 1", "year 1": "Year 1", "Year 1": "Year 1",
        "3": "Year 3", "year 3": "Year 3", "Year 3": "Year 3",
        "5": "Year 5", "year 5": "Year 5", "Year 5": "Year 5",
    }
    normalized_year = year_map.get(str(year).strip(), year)
    
    # Path to the assessment criteria reference
    skill_path = Path(__file__).parent / "skills" / "myp-rubric-creator" / "references" / "assessment_criteria.md"
    
    if not skill_path.exists():
        return f"Error: Assessment criteria file not found at {skill_path}"
    
    content = skill_path.read_text(encoding="utf-8")
    
    # Parse and extract the relevant section
    # Find the year section
    year_header = f"# {normalized_year}"
    year_start = content.find(year_header)
    if year_start == -1:
        return f"Error: Could not find {normalized_year} in assessment criteria."
    
    # Find the criterion section within the year
    criterion_header = f"## Criterion {criterion}:"
    criterion_start = content.find(criterion_header, year_start)
    if criterion_start == -1:
        return f"Error: Could not find Criterion {criterion} in {normalized_year}."
    
    # Find the end of this criterion section (next ## or #)
    next_section = content.find("\n##", criterion_start + len(criterion_header))
    if next_section == -1:
        section_content = content[criterion_start:]
    else:
        section_content = content[criterion_start:next_section]
    
    return section_content.strip()


def get_year_level_command_terms(
    year: Literal["Year 1", "Year 3", "Year 5", "1", "3", "5"],
    criterion: Literal["A", "B", "C", "D"] = "C"
) -> list[str]:
    """
    Extract command terms from assessment criteria for a specific year and criterion.
    
    Reads the assessment_criteria.md file and extracts all command terms
    marked with **bold** formatting from the specified year and criterion section.
    
    Args:
        year: MYP year level (Year 1, Year 3, or Year 5)
        criterion: Assessment criterion (A, B, C, or D). Default is C.
    
    Returns:
        List of unique command terms extracted from the rubric descriptors.
    """
    import re
    
    year_map = {
        "1": "Year 1", "year 1": "Year 1", "Year 1": "Year 1",
        "3": "Year 3", "year 3": "Year 3", "Year 3": "Year 3",
        "5": "Year 5", "year 5": "Year 5", "Year 5": "Year 5",
    }
    normalized_year = year_map.get(str(year).strip(), year)
    
    # Read the assessment criteria file
    project_root = Path(__file__).parent
    skills_path = project_root / "skills" / "myp-rubric-creator" / "references" / "assessment_criteria.md"
    
    try:
        with open(skills_path.resolve(), 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        return [f"Error reading assessment criteria: {str(e)}"]
    
    # Find the section for the specified year and criterion
    # Pattern: ## Criterion X: ... under # Year X section
    year_pattern = rf"# {normalized_year}\s*\n(.*?)# Year"
    year_match = re.search(year_pattern, content, re.DOTALL | re.IGNORECASE)
    
    if not year_match:
        # Try to match until end of file for Year 5
        year_pattern = rf"# {normalized_year}\s*\n(.*)"
        year_match = re.search(year_pattern, content, re.DOTALL | re.IGNORECASE)
    
    if not year_match:
        return [f"Year section '{normalized_year}' not found"]
    
    year_section = year_match.group(1)
    
    # Find the criterion section within the year section
    criterion_pattern = rf"## Criterion {criterion}:.*?(?=## Criterion|# Year|$)"
    criterion_match = re.search(criterion_pattern, year_section, re.DOTALL | re.IGNORECASE)
    
    if not criterion_match:
        return [f"Criterion {criterion} not found in {normalized_year}"]
    
    criterion_section = criterion_match.group(0)
    
    # Extract all **bold** terms (command terms are bolded in the rubric)
    # Pattern: **word** or **word word**
    bold_pattern = r"\*\*([a-zA-Z\s]+?)\*\*"
    matches = re.findall(bold_pattern, criterion_section)
    
    # Clean up and deduplicate
    command_terms = []
    seen = set()
    
    for term in matches:
        # Clean up whitespace and convert to lowercase
        clean_term = term.strip().lower()
        # Skip empty or very short terms
        if len(clean_term) > 2 and clean_term not in seen:
            # Skip common non-command words
            if clean_term not in ['the', 'and', 'for', 'maximum', 'year', 'criterion']:
                command_terms.append(clean_term)
                seen.add(clean_term)
    
    return command_terms


# =============================================================================
# MEMORY MANAGEMENT TOOLS
# =============================================================================

# Global reference to memory store (set by agent during initialization)
_memory_store = None


def set_memory_store(store):
    """Set the global memory store reference. Called by agent during initialization."""
    global _memory_store
    _memory_store = store


def _get_memory_store():
    """Get the memory store, raising error if not set."""
    if _memory_store is None:
        raise RuntimeError("Memory store not initialized. Call set_memory_store() first.")
    return _memory_store


def read_memory() -> str:
    """
    Read the current user memory from the store.
    
    Use this to check what you already know about the user.
    
    Returns:
        The current user memory as a formatted string, or a message if no memory exists.
    
    Example:
        read_memory()
        # Returns: "- User is a PE teacher at International School\n- Teaches MYP Year 3..."
    """
    try:
        # Import here to avoid circular imports
        from memory import load_memory_from_store, format_memory_for_prompt, get_memory_namespace, MEMORY_KEY
        
        store = _get_memory_store()
        
        # Use a default user ID (same as chat.py uses)
        # In practice, this should be passed through config
        user_id = "default"
        
        memory_data = load_memory_from_store(store, user_id)
        
        if not memory_data or not memory_data.get("memory"):
            return "No memory stored yet. This is a fresh conversation."
        
        return f"Current user memory:\n\n{memory_data.get('memory')}"
        
    except Exception as e:
        return f"Error reading memory: {str(e)}"


def update_memory(fact: str) -> str:
    """
    Add a specific fact to the user memory.
    
    Use this when the user explicitly tells you something important
    that should be remembered for future conversations.
    
    Args:
        fact: The fact to remember about the user.
              Example: "User teaches at Singapore American School"
              Example: "User prefers detailed rubrics with examples"
    
    Returns:
        Confirmation that the fact was added to memory.
    
    Example:
        update_memory("User is an IB coordinator for MYP PHE")
    """
    try:
        from memory import load_memory_from_store, save_memory_to_store, get_memory_namespace, MEMORY_KEY
        from datetime import datetime
        
        store = _get_memory_store()
        user_id = "default"
        
        # Load existing memory
        memory_data = load_memory_from_store(store, user_id)
        
        # Get current memory content
        if memory_data and memory_data.get("memory"):
            current_memory = memory_data.get("memory")
        else:
            current_memory = ""
        
        # Append new fact
        if current_memory:
            new_memory = current_memory + f"\n- {fact}"
        else:
            new_memory = f"- {fact}"
        
        # Save updated memory
        updated_data = {
            "memory": new_memory,
            "last_updated": datetime.now().isoformat()
        }
        save_memory_to_store(store, user_id, updated_data)
        
        return f"✅ Added to memory: {fact}"
        
    except Exception as e:
        return f"Error updating memory: {str(e)}"


def clear_memory(confirm: bool = False) -> str:
    """
    Clear all user memory from the store.
    
    WARNING: This will permanently delete all learned information about the user.
    
    Args:
        confirm: Must be set to True to confirm deletion.
    
    Returns:
        Confirmation of memory clearance or warning if not confirmed.
    
    Example:
        clear_memory(confirm=True)
    """
    if not confirm:
        return "⚠️  Warning: This will delete all user memory. Set confirm=True to proceed."
    
    try:
        from memory import get_memory_namespace, MEMORY_KEY
        
        store = _get_memory_store()
        user_id = "default"
        
        namespace = get_memory_namespace(user_id)
        store.put(namespace, MEMORY_KEY, {"memory": "", "last_updated": None})
        
        return "🗑️  All user memory has been cleared."
        
    except Exception as e:
        return f"Error clearing memory: {str(e)}"


# =============================================================================
# EXA SEARCH TOOL
# =============================================================================

def search_exa(
    query: str,
    search_type: Literal["auto", "instant", "fast", "deep"] = "auto",
    num_results: int = 5,
    include_highlights: bool = True,
    include_summary: bool = False,
    category: Optional[Literal["company", "people", "research paper", "news", "personal site", "financial report"]] = None
) -> dict:
    """
    Search the web using Exa AI search engine.
    
    Exa is a custom search engine built for AIs that offers:
    - Neural search with semantic understanding
    - Structured content extraction (highlights, summaries)
    - Custom data types (companies, people, research papers)
    
    Args:
        query: The search query (natural language)
        search_type: Search speed/quality tradeoff
            - "auto": ~1s, default balanced option
            - "instant": ~200ms, for real-time apps
            - "fast": ~450ms, speed with quality
            - "deep": 5-60s, complex multi-step reasoning
        num_results: Number of results (1-10)
        include_highlights: Return token-efficient content highlights
        include_summary: Return AI-generated summaries
        category: Filter by data type (company, people, research paper, etc.)
    
    Returns:
        Dictionary with search results containing titles, URLs, and content.
    
    Example:
        search_exa("latest research on physical education assessment", type="deep")
    """
    api_key = os.getenv("EXA_API_KEY")
    if not api_key:
        return {
            "error": "EXA_API_KEY not found. Please set it in your .env file.",
            "signup_url": "https://dashboard.exa.ai/"
        }
    
    # Build contents configuration
    contents = {}
    if include_highlights:
        contents["highlights"] = {"max_characters": 4000}
    if include_summary:
        contents["summary"] = True
    
    # Build request payload
    payload = {
        "query": query,
        "type": search_type,
        "numResults": min(max(num_results, 1), 10),
        "contents": contents if contents else {"text": True}
    }
    
    if category:
        payload["category"] = category
    
    try:
        response = requests.post(
            "https://api.exa.ai/search",
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key
            },
            json=payload,
            timeout=120 if search_type == "deep" else 30
        )
        response.raise_for_status()
        
        data = response.json()
        
        # Format results for readability
        formatted_results = []
        for result in data.get("results", []):
            formatted_result = {
                "title": result.get("title", "No title"),
                "url": result.get("url", ""),
                "published_date": result.get("publishedDate", "Unknown"),
                "author": result.get("author", "Unknown"),
            }
            
            if include_highlights and "highlights" in result:
                formatted_result["highlights"] = result["highlights"]
            
            if include_summary and "summary" in result:
                formatted_result["summary"] = result["summary"]
            
            if "text" in result:
                formatted_result["text_preview"] = result["text"][:500] + "..." if len(result["text"]) > 500 else result["text"]
            
            formatted_results.append(formatted_result)
        
        return {
            "success": True,
            "query": query,
            "search_type": search_type,
            "num_results": len(formatted_results),
            "results": formatted_results,
            "cost": data.get("costDollars", {}).get("total", "unknown")
        }
        
    except requests.exceptions.RequestException as e:
        return {
            "error": f"Exa API request failed: {str(e)}",
            "query": query
        }


def search_exa_structured(
    query: str,
    output_schema: dict,
    num_results: int = 5
) -> dict:
    """
    Search using Exa with structured JSON output extraction.
    
    Uses deep search to extract structured data from web results.
    Perfect for data enrichment and research tasks.
    
    Args:
        query: The search query
        output_schema: JSON schema defining the structure to extract
        num_results: Number of results to process
    
    Returns:
        Structured data extracted from search results.
    
    Example:
        search_exa_structured(
            "top sports education companies",
            output_schema={
                "type": "object",
                "properties": {
                    "companies": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "name": {"type": "string"},
                                "focus_area": {"type": "string"}
                            }
                        }
                    }
                }
            }
        )
    """
    api_key = os.getenv("EXA_API_KEY")
    if not api_key:
        return {
            "error": "EXA_API_KEY not found. Please set it in your .env file.",
            "signup_url": "https://dashboard.exa.ai/"
        }
    
    payload = {
        "query": query,
        "type": "deep",
        "numResults": min(max(num_results, 1), 5),
        "outputSchema": output_schema,
        "contents": {"text": {"max_characters": 20000}}
    }
    
    try:
        response = requests.post(
            "https://api.exa.ai/search",
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key
            },
            json=payload,
            timeout=120
        )
        response.raise_for_status()
        
        data = response.json()
        
        return {
            "success": True,
            "query": query,
            "structured_output": data.get("output", {}).get("content", {}),
            "grounding": data.get("output", {}).get("grounding", []),
            "results": [
                {
                    "title": r.get("title"),
                    "url": r.get("url"),
                    "author": r.get("author")
                }
                for r in data.get("results", [])
            ]
        }
        
    except requests.exceptions.RequestException as e:
        return {
            "error": f"Exa API request failed: {str(e)}",
            "query": query
        }



def _resolve_workspace_path(path: str) -> Path:
    """
    Resolve a user-provided path relative to the workspace root.
    
    This follows the same logic as FilesystemBackend._resolve_path() in virtual_mode:
    - Virtual paths starting with '/' are resolved under WORKSPACE_ROOT
    - Relative paths are also resolved under WORKSPACE_ROOT
    - Path traversal ('..', '~') is blocked
    
    Args:
        path: The user-provided path string (e.g., '/drafts/file.docx' or 'drafts/file.docx')
        
    Returns:
        Resolved absolute Path object
        
    Raises:
        ValueError: If path traversal is attempted
    """
    # Block path traversal attempts (same as FilesystemBackend)
    if ".." in path or path.startswith("~"):
        raise ValueError("Path traversal not allowed")
    
    # Normalize to virtual path (ensure it starts with /)
    vpath = path if path.startswith("/") else "/" + path
    
    # Resolve under workspace root (strip leading / and join with WORKSPACE_ROOT)
    full_path = (WORKSPACE_ROOT / vpath.lstrip("/")).resolve()
    
    # Ensure the resolved path stays within WORKSPACE_ROOT (security check)
    try:
        full_path.relative_to(WORKSPACE_ROOT.resolve())
    except ValueError as e:
        raise ValueError(f"Path outside workspace directory: {path}") from e
    
    return full_path


def read_docx(path: str) -> str:
    """
    Read and extract text from a Microsoft Word (.docx) file.
    
    Use this tool to read .docx files. It extracts text from paragraphs
    and tables, preserving the document structure.
    
    Args:
        path: Path to the .docx file, relative to workspace root.
              Can start with or without leading '/'.
              Virtual paths like '/drafts/file.docx' are resolved under workspace/.
              Examples: "drafts/document.docx"
                       "/drafts/document.docx" (both resolve to workspace/drafts/)
    
    Returns:
        The extracted text content from the Word document.
    
    Example:
        read_docx("drafts/rubric.docx")
        read_docx("/drafts/rubric.docx")
    """
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx library not installed. Run: pip install python-docx"
    
    try:
        # Resolve the path using workspace-relative logic (matches FilesystemBackend)
        file_path = _resolve_workspace_path(path)
        
        # Check if file exists
        if not file_path.exists():
            return f"Error: File not found: {path} (tried: {file_path})"
        
        # Check if it's a file
        if not file_path.is_file():
            return f"Error: Path is not a file: {path}"
        
        # Check file extension
        if file_path.suffix.lower() != '.docx':
            return f"Error: File is not a .docx file: {path}"
        
        # Load and parse the document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        tables = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(row_text))
            if table_rows:
                tables.append("\n".join(table_rows))
        
        # Combine content
        content_parts = []
        if paragraphs:
            content_parts.append("\n\n".join(paragraphs))
        if tables:
            content_parts.append("\n\n--- Tables ---\n\n" + "\n\n".join(tables))
        
        result = "\n\n".join(content_parts) if content_parts else "(Document appears to be empty)"
        
        return result
        
    except PermissionError:
        return f"Error: Permission denied reading file: {path}"
    except Exception as e:
        return f"Error reading .docx file {path}: {str(e)}"
